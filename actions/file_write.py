"""file_write — JARVIS'in dosya oluşturma/yazma yetisi.

"Bir özet yaz ve rapor.md olarak kaydet", "Şu listeyi notlar.txt'e ekle"
gibi istekleri karşılar. Metin, Markdown ve Word (.docx) destekler.
"""
from __future__ import annotations

import os
from pathlib import Path

_DEFAULT_DIR = Path.home() / "Desktop"
_MAX_BYTES = 10_000_000   # 10 MB yazma tavanı


def _resolve_path(target: str) -> Path:
    p = Path(os.path.expanduser(os.path.expandvars(target)))
    if not p.is_absolute():
        p = _DEFAULT_DIR / p
    return p


def _ensure_parent(p: Path) -> None:
    p.parent.mkdir(parents=True, exist_ok=True)


def _write_text(p: Path, content: str, append: bool) -> str:
    mode = "a" if append else "w"
    with open(p, mode, encoding="utf-8", newline="") as f:
        f.write(content)
    return f"{'Eklendi' if append else 'Yazıldı'}: {p} ({len(content)} karakter)."


def _write_docx(p: Path, content: str, title: str | None) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Hata: python-docx kurulu değil. `pip install python-docx` çalıştır."
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.save(str(p))
    return f"Word dosyası yazıldı: {p}."


def file_write_action(parameters: dict | None = None, player=None) -> str:
    params = parameters or {}
    target = (params.get("path") or params.get("target") or "").strip()
    content = params.get("content") or ""
    fmt = (params.get("format") or "").strip().lower()
    title = params.get("title")
    append = bool(params.get("append"))
    overwrite = bool(params.get("overwrite"))

    if not target:
        return "Hata: path parametresi gerekli."
    if not isinstance(content, str):
        content = str(content)
    if len(content.encode("utf-8", errors="ignore")) > _MAX_BYTES:
        return "Hata: içerik 10 MB sınırını aşıyor."

    path = _resolve_path(target)
    if not fmt:
        fmt = path.suffix.lower().lstrip(".") or "txt"

    if path.exists() and not append and not overwrite:
        return (
            f"Hata: '{path}' zaten var. overwrite=true ya da append=true ile çağır."
        )

    try:
        _ensure_parent(path)
        if fmt in ("txt", "md", "markdown", "log", "csv", "json", "html", "py"):
            return _write_text(path, content, append)
        if fmt in ("docx", "doc"):
            return _write_docx(path, content, title)
        return _write_text(path, content, append)
    except PermissionError:
        return f"Hata: '{path}' üzerine yazma izni yok."
    except Exception as e:
        return f"Hata: dosya yazılamadı ({type(e).__name__}: {e})."
